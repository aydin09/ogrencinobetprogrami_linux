from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.shared import Length
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk
import calendar
import locale
from calendar import weekday, day_name
from itertools import cycle

locale.setlocale(locale.LC_ALL, '')

def gunler(event):
    global gun,ay,yil
    ay=aylar.get(ACTIVE)
    ttk.Label(mainframe, text ="                        ").grid(column = 6, row=2)
    ttk.Label(mainframe, text ="YILI GİRİNİZ").grid(column = 6, row=0)
    yil = ttk.Entry(mainframe, width =30)
    yil.grid(column = 6, row = 1)
    
    ttk.Label(mainframe, text =ay).grid(column = 6, row=2)
    ttk.Label(mainframe, text ='AYININ TATİL GÜNLERİNİ ARALARINA VİRGÜL KOYARAK SAYI OLARAK GİRİNİZ').grid(column = 6, row=3)
    gun = ttk.Entry(mainframe, width =30)
    gun.grid(column = 6, row = 4)
    gun.insert(0,0)
    ttk.Button(mainframe, text='Çıktı Al',command= cikti).grid(column=6, row=5)

def bilgi_girişi(event):
    liste1=liste.get(ACTIVE)

    ogrenci.delete(0,END)

    ogrenci.insert(END,liste1)
                        
def nobet():
    okul_adi1 = okul_adi.get()
    sinif1 = sinif.get()
    sube1 = sube.get()

    if okul_adi1 =="" or sinif1 =="" or sube1 =="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='nobet.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Bilgileri tam giriniz!').pack()

    else:    
        okul_adi.delete(0,END)
        sinif.delete(0,END)
        sube.delete(0,END)

        vt1 = sqlite3.connect('nobet.sql')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS nobet(okuladi TEXT, sinif TEXT, sube TEXT)""")
        im1.execute("""DELETE FROM nobet""")
        im1.execute("""INSERT INTO nobet VALUES  (?,?,?)""",(okul_adi1,sinif1,sube1,))

        im1.execute("""SELECT * FROM  nobet""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}{}"
        for row1 in rows:
            data_str += sf.format(row1[0],row1[1],row1[2])

            okul_adi.insert(END,row1[0])
            sinif.insert(END,row1[1])
            sube.insert(END,row1[2])

        vt1.commit()

def kaydet_ogrenci():
    ogrenci1 = ogrenci.get()

    if ogrenci1 =="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='nobet.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Öğrencinin adını soyadını giriniz!').pack()

    else:
        ogrenci.delete(0,END)

        liste.delete(0,END)
        
        vt2 = sqlite3.connect('ogrenci.sql3')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS ogrenci(ogrenci TEXT)""")
        im2.execute("""INSERT INTO ogrenci VALUES  (?)""",(ogrenci1,))
        im2.execute("""SELECT * FROM  ogrenci""")
        rows2 = im2.fetchall()
        data_str2 = ""
        sf2 = "{}"
        for row2 in rows2:
            data_str2 += sf2.format(row2[0])

            liste.insert(END,row2[0])

        vt2.commit()

def cikti():
    yil2=""
    gun1=gun.get()
    yil1=yil.get()
    yil2+=yil1

    vt2 = sqlite3.connect('ogrenci.sql3')
    im2= vt2.cursor()
    im2.execute("""CREATE TABLE IF NOT EXISTS ogrenci(ogrenci TEXT)""")
    im2.execute("""SELECT * FROM  ogrenci""")
    rows2 = im2.fetchall()
    data_str2 = []
    sf2 = "{}"
    for row2 in rows2:
        data_str2.append(sf2.format(row2[0]))

    vt2.commit()
    
    if yil1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='nobet.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Yılı giriniz!').pack()

    else:
        gun_liste=""
        gun_liste +=gun1
        gun_sayi=gun_liste.split(",")

        vt1 = sqlite3.connect('nobet.sql')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS nobet(okuladi TEXT, sinif TEXT, sube TEXT)""")
        im1.execute("""SELECT * FROM  nobet""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}{}"
        for row1 in rows:
            data_str += sf.format(row1[0],row1[1],row1[2])

        vt1.commit()
          
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)    

        paragraph = document.add_paragraph()
        paragraph.add_run(row1[0]+" "+row1[1]+"/"+row1[2]+" SINIFI").bold = True
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.alignment = 1

        paragraph = document.add_paragraph()
        paragraph.add_run(ay+" AYI NÖBET LİSTESİ").bold = True
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.alignment = 1

        lis=cycle(data_str2)

        if ay == "EYLÜL":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30]
            table = document.add_table(rows=31, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=31
             
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 9, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 9, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 9, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Eylül ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "EKİM":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]
            table = document.add_table(rows=32, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=32
                       
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 10, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 10, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 10, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Ekim ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "KASIM":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30]
            table = document.add_table(rows=31, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=31
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 11, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 11, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 11, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Kasım ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "ARALIK":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]
            table = document.add_table(rows=32, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=32
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 12, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 12, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 12, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Aralık ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "OCAK":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]
            table = document.add_table(rows=32, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=32
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 1, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 1, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 1, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Ocak ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "ŞUBAT":
            if int(yil2) % 4 == 0 and int(yil2) % 100 != 0:
                liste=set()
                fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29]
                table = document.add_table(rows=30, cols=2,style = 'Table Grid')
                cell = table.cell(0,0)
                table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
                table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[0].width = Cm(7)

                cell = table.cell(0,1)
                table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
                table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[1].width = Cm(10)

                a=30
           
                for s in range(1,a):
                    cell = table.cell(s,0)
                    cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 2, s)]
                    table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for g in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g)]=="Cumartesi":
                        cell = table.cell(int(g),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g)

                for g1 in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g1)]=="Pazar":
                        cell = table.cell(int(g1),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g1)

                if gun_sayi[0] == "0":
                    uyari=Toplevel()
                    uyari.resizable(width=FALSE ,height=FALSE)
                    img=PhotoImage(file='nobet.png')
                    uyari.tk.call('wm','iconphoto',uyari._w,img)
                    Label(uyari, text ='Şubat ayında tatil yok.').pack()
                else:
                    for i in gun_sayi:
                        cell = table.cell(int(i),1)
                        cell.text ="TATİL"
                        table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(int(i))

                for i in liste:
                    fark.remove(i)

                for o in fark:
                    cell = table.cell(int(o),1)
                    cell.text =next(lis)+" - "+next(lis)
                    table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            elif int(yil2) % 100 == 0:
                liste=set()
                fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28]
                table = document.add_table(rows=29, cols=2,style = 'Table Grid')
                cell = table.cell(0,0)
                table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
                table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[0].width = Cm(7)

                cell = table.cell(0,1)
                table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
                table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[1].width = Cm(10)

                a=29
           
                for s in range(1,a):
                    cell = table.cell(s,0)
                    cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 2, s)]
                    table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for g in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g)]=="Cumartesi":
                        cell = table.cell(int(g),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g)

                for g1 in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g1)]=="Pazar":
                        cell = table.cell(int(g1),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g1)

                if gun_sayi[0] == "0":
                    uyari=Toplevel()
                    uyari.resizable(width=FALSE ,height=FALSE)
                    img=PhotoImage(file='nobet.png')
                    uyari.tk.call('wm','iconphoto',uyari._w,img)
                    Label(uyari, text ='Şubat ayında tatil yok.').pack()
                else:
                    for i in gun_sayi:
                        cell = table.cell(int(i),1)
                        cell.text ="TATİL"
                        table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(int(i))

                for i in liste:
                    fark.remove(i)

                for o in fark:
                    cell = table.cell(int(o),1)
                    cell.text =next(lis)+" - "+next(lis)
                    table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            elif int(yil2) % 400 ==0:
                liste=set()
                fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29]
                table = document.add_table(rows=30, cols=2,style = 'Table Grid')
                cell = table.cell(0,0)
                table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
                table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[0].width = Cm(7)

                cell = table.cell(0,1)
                table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
                table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[1].width = Cm(10)

                a=30
        
                for s in range(1,a):
                    cell = table.cell(s,0)
                    cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 2, s)]
                    table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for g in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g)]=="Cumartesi":
                        cell = table.cell(int(g),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g)

                for g1 in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g1)]=="Pazar":
                        cell = table.cell(int(g1),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g1)

                if gun_sayi[0] == "0":
                    uyari=Toplevel()
                    uyari.resizable(width=FALSE ,height=FALSE)
                    img=PhotoImage(file='nobet.png')
                    uyari.tk.call('wm','iconphoto',uyari._w,img)
                    Label(uyari, text ='Şubat ayında tatil yok.').pack()
                else:
                    for i in gun_sayi:
                        cell = table.cell(int(i),1)
                        cell.text ="TATİL"
                        table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(int(i))

                for i in liste:
                    fark.remove(i)

                for o in fark:
                    cell = table.cell(int(o),1)
                    cell.text =next(lis)+" - "+next(lis)
                    table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            else:
                liste=set()
                fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28]
                table = document.add_table(rows=29, cols=2,style = 'Table Grid')
                cell = table.cell(0,0)
                table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
                table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[0].width = Cm(7)

                cell = table.cell(0,1)
                table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
                table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.columns[1].width = Cm(10)

                a=29
           
                for s in range(1,a):
                    cell = table.cell(s,0)
                    cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 2, s)]
                    table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

                for g in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g)]=="Cumartesi":
                        cell = table.cell(int(g),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g)

                for g1 in range(1,a):
                    if calendar.day_name[calendar.weekday(int(yil2), 2, g1)]=="Pazar":
                        cell = table.cell(int(g1),1)
                        cell.text ="HAFTA SONU"
                        table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(g1)

                if gun_sayi[0] == "0":
                    uyari=Toplevel()
                    uyari.resizable(width=FALSE ,height=FALSE)
                    img=PhotoImage(file='nobet.png')
                    uyari.tk.call('wm','iconphoto',uyari._w,img)
                    Label(uyari, text ='Şubat ayında tatil yok.').pack()
                else:
                    for i in gun_sayi:
                        cell = table.cell(int(i),1)
                        cell.text ="TATİL"
                        table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                        liste.add(int(i))

                for i in liste:
                    fark.remove(i)

                for o in fark:
                    cell = table.cell(int(o),1)
                    cell.text =next(lis)+" - "+next(lis)
                    table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "MART":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]
            table = document.add_table(rows=32, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=32
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 3, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 3, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 3, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Mart ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "NİSAN":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30]
            table = document.add_table(rows=31, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=31
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 4, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 4, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 4, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Nisan ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        elif ay == "MAYIS":
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]
            table = document.add_table(rows=32, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=32
                           
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 5, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 5, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 5, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Mayıs ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        else:
            liste=set()
            fark=[1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30]
            table = document.add_table(rows=31, cols=2,style = 'Table Grid')
            cell = table.cell(0,0)
            table.cell(0,0).paragraphs[0].add_run("GÜNLER").bold = True
            table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[0].width = Cm(7)

            cell = table.cell(0,1)
            table.cell(0,1).paragraphs[0].add_run("NÖBETÇİLER").bold = True
            table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.columns[1].width = Cm(10)

            a=31
                            
            for s in range(1,a):
                cell = table.cell(s,0)
                cell.text =str(s)+"/"+ay+"/"+str(yil2)+" - " +calendar.day_name[calendar.weekday(int(yil2), 6, s)]
                table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            for g in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 6, g)]=="Cumartesi":
                    cell = table.cell(int(g),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g)

            for g1 in range(1,a):
                if calendar.day_name[calendar.weekday(int(yil2), 6, g1)]=="Pazar":
                    cell = table.cell(int(g1),1)
                    cell.text ="HAFTA SONU"
                    table.cell(int(g1),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(g1)

            if gun_sayi[0] == "0":
                uyari=Toplevel()
                uyari.resizable(width=FALSE ,height=FALSE)
                img=PhotoImage(file='nobet.png')
                uyari.tk.call('wm','iconphoto',uyari._w,img)
                Label(uyari, text ='Haziran ayında tatil yok.').pack()
            else:
                for i in gun_sayi:
                    cell = table.cell(int(i),1)
                    cell.text ="TATİL"
                    table.cell(int(i),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                    liste.add(int(i))

            for i in liste:
                fark.remove(i)

            for o in fark:
                cell = table.cell(int(o),1)
                cell.text =next(lis)+" - "+next(lis)
                table.cell(int(o),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
       
        document.save('nobet.docx')

        os.system("libreoffice --writer nobet.docx")

def sil_ogrenci():
    ogrenci1 = ogrenci.get()

    if ogrenci1 =="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='nobet.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Listeden öğrenciyi çift tıklayarak seçiniz. Sonra Sil butonuna basınız.').pack()

    else:
        liste.delete(0,END)

        ogrenci.delete(0,END)

        vt2 = sqlite3.connect('ogrenci.sql3')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS ogrenci(ogrenci TEXT)""")
        im2.execute("""DELETE FROM ogrenci WHERE ogrenci =?""",(ogrenci1,))
        im2.execute("""SELECT * FROM  ogrenci""")
        rows2 = im2.fetchall()
        data_str2 = ""
        sf2 = "{}"
        for row2 in rows2:
            data_str2 += sf2.format(row2[0])

            liste.insert(END,row2[0])

        vt2.commit()

root = Tk()
root.title("ÖĞRENCİ NÖBET PROGRAMI")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='nobet.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

okul_adi = ttk.Entry(mainframe, width =30)
okul_adi.grid(column = 2, row = 0)

sinif = ttk.Entry(mainframe, width =30)
sinif.grid(column = 2, row = 1)

sube = ttk.Entry(mainframe, width =30)
sube.grid(column = 2, row = 2)

ogrenci = ttk.Entry(mainframe, width =30)
ogrenci.grid(column = 3, row = 34)

ttk.Label(mainframe, text ='OKULUN ADI').grid(column = 1, row = 0, sticky=W)
ttk.Label(mainframe, text ='SINIFI').grid(column = 1, row = 1, sticky=W)
ttk.Label(mainframe, text ='ŞUBESİ').grid(column = 1, row=2, sticky=W)

ttk.Label(mainframe, text ='ÖĞRENCİ LİSTESİ').grid(column = 3, row=0)
ttk.Label(mainframe, text ='ÖĞRENCİNİN ADI SOYADI').grid(column = 3, row=33)

ttk.Label(mainframe, text ='AYI ÇİFT TIKLAYIP SEÇİNİZ').grid(column = 5, row=0)

aylar = Listbox(mainframe,width=33)
aylar.grid(column=5, row=1,rowspan=33,  sticky=(N,S,E,W))

ayAdlari = ["EYLÜL","EKİM","KASIM","ARALIK","OCAK","ŞUBAT","MART","NİSAN","MAYIS","HAZİRAN"]
for i in ayAdlari:
    aylar.insert(END, i)

aylar.bind("<Double-Button-1>", gunler)

ttk.Label(mainframe, text ='').grid(column = 3, row=32)
ttk.Label(mainframe, text ='').grid(column = 3, row=35)
ttk.Label(mainframe, text ='').grid(column = 3, row=37)

liste = Listbox(mainframe,width=32)
liste.grid(column=3, row=1,rowspan=32,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",bilgi_girişi)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=4, row=1, rowspan=32,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

vt1 = sqlite3.connect('nobet.sql')
im1= vt1.cursor()
im1.execute("""CREATE TABLE IF NOT EXISTS nobet(okuladi TEXT, sinif TEXT, sube TEXT)""")
im1.execute("""SELECT * FROM  nobet""")
rows = im1.fetchall()
data_str = ""
sf = "{}{}{}"
for row1 in rows:
    data_str += sf.format(row1[0],row1[1],row1[2])

    okul_adi.insert(END,row1[0])
    sinif.insert(END,row1[1])
    sube.insert(END,row1[2])

vt1.commit()

vt2 = sqlite3.connect('ogrenci.sql3')
im2= vt2.cursor()
im2.execute("""CREATE TABLE IF NOT EXISTS ogrenci(ogrenci TEXT)""")
im2.execute("""SELECT * FROM  ogrenci""")
rows2 = im2.fetchall()
data_str2 = ""
sf2 = "{}"
for row2 in rows2:
    data_str2 += sf2.format(row2[0])

    liste.insert(END,row2[0])

vt2.commit()

ttk.Button(mainframe, text='Kaydet',command= nobet).grid(column=2, row=4)
ttk.Button(mainframe, text='Öğrenci Kaydet',command= kaydet_ogrenci).grid(column=3, row=36)
ttk.Button(mainframe, text='Sil', command= sil_ogrenci).grid(column=3, row=38)

okul_adi.focus()

root.mainloop()    
